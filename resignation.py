from docx import Document

# Create a new Document
doc = Document()

# Add a Title to the document
doc.add_heading('Resignation Letter', 0)

# Add the content
doc.add_paragraph("""
Yuvraj Pratap Singh
[Your Address]
[City, State, Zip Code]
[Email Address]
[Phone Number]
[Date]

The chairperson 
ACM Students Chapter
[University/Organization Name]
[University Address]
[City, State, Zip Code]

Dear [President's Name],

I hope this letter finds you well. I am writing to formally resign from my position as [Your Position] of the ACM Students Chapter at [University/Organization Name], effective [Last Working Day, e.g., two weeks from the date of this letter].

This decision has not been an easy one, as my time with the ACM Students Chapter has been incredibly rewarding and has provided me with numerous opportunities to grow both personally and professionally. However, after careful consideration, I have decided that it is time for me to step down and focus on [reason for resigning, e.g., academic commitments, personal projects, or other responsibilities].

I am deeply grateful for the experiences and friendships I have gained through my involvement with the chapter. It has been a pleasure working alongside such a talented and dedicated team, and I am proud of the accomplishments we have achieved together.

I will do my utmost to ensure a smooth transition during this period. Please let me know how I can assist in this process, whether it be by helping to train my successor or providing any other support that may be needed.

Thank you for the trust and support you have extended to me during my tenure. I look forward to staying connected with the chapter and contributing in any way I can in the future.

Wishing the ACM Students Chapter continued success.

Sincerely,

Yuvraj Pratap Singh
""")

# Save the document
file_path = "/mnt/data/Yuvraj_Pratap_Singh_Resignation_Letter.docx"
doc.save(file_path)

file_path
